from flask import Flask, request, jsonify, send_file
from openai import OpenAI
from flask_cors import CORS
import docx
import os

app = Flask(__name__)

CORS(app, resources={r"/generer_fiche_cours": {"origins": "*"}})

# Clé API OpenAI
openai_api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=openai_api_key)

# Fonction pour interroger l'API OpenAI avec les données et générer du contenu
def generer_fiche_de_cours(template, document_data):
    # Créer des messages pour l'API chat
    messages = [
        {"role": "system", "content": "Tu es un assistant qui aide à créer des fiches de cours en suivant un modèle donné."},
        {"role": "user", "content": f"Voici un modèle de fiche de cours : {template}"},
        {"role": "user", "content": f"Voici les informations à utiliser pour remplir ce modèle : {document_data}"}
        {"role": "user", "content": f"Si la durée du cours n'est pas spécifier alors tu dois l'estimer."},
    ]

    # Utilisation de l'endpoint chat completions avec un modèle de type chat
    response = client.chat.completions.create(
        model="gpt-4",  # Ou "gpt-4" si disponible
        messages=messages,
        max_tokens=1000,
        temperature=0.7
    )

    return response.choices[0].message.content.strip()

# Fonction pour créer un fichier Word à partir du contenu généré
def creer_document_word(titre, contenu):
    doc = docx.Document()

    # Ajouter le titre
    doc.add_heading(titre, 0)

    # Ajouter le contenu
    doc.add_paragraph(contenu)

    # Sauvegarder le fichier Word
    nom_fichier = f"{titre}.docx"
    doc.save(nom_fichier)

    return nom_fichier

# API endpoint pour générer la fiche de cours
@app.route('/generer_fiche_cours', methods=['POST'])
def generer_fiche_cours_api():
    data = request.get_json()

    template = data.get('template', '')
    document_data = data.get('document_data', '')
    titre = data.get('titre', 'Fiche de cours')

    if not template or not document_data:
        return jsonify({'error': 'Le template et les données du document sont requis'}), 400

    # Générer la fiche de cours avec OpenAI
    fiche_de_cours = generer_fiche_de_cours(template, document_data)

    # Créer un fichier Word avec la fiche de cours générée
    fichier_word = creer_document_word(titre, fiche_de_cours)

    # Envoyer le fichier Word en tant que réponse
    return send_file(fichier_word, as_attachment=True)

# Démarrer le serveur Flask
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
